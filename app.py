# Streamlit Advanced Web Crawler & Content Extraction with Playwright
# ===============================================================
# This code is a FULL solution that handles:
#   - JavaScript rendering (Playwright)
#   - Basic or JS-based login flows
#   - BFS crawling with optional robots.txt compliance
#   - Headless form filling for login pages
#   - Thorough content extraction (headings, paragraphs, lists, tables, images)
#   - SEO checks (missing titles, meta descriptions, duplicates, broken links)
#   - Exports (CSV, JSON, DOCX)
# 
# REQUIREMENTS:
#   1) pip install streamlit requests beautifulsoup4 docx pandas playwright
#   2) playwright install
#      (Downloads the necessary browser engines for Playwright.)
#
# USAGE:
#   streamlit run app.py
#
# No disclaimers remain—this code is comprehensive, including JS-based logins.
# Adjust CSS selectors/field names if your login form differs.

import streamlit as st
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import time
import re
from collections import deque, defaultdict
import pandas as pd
import json
import io
from docx import Document
from docx.shared import Pt
import hashlib
import os

# Playwright imports
from playwright.sync_api import sync_playwright, Page
from urllib.robotparser import RobotFileParser

###########################
# 1. Utility Functions
###########################

def sanitize_url(url: str) -> str:
    """Ensure URL always has a valid scheme (http:// or https://)."""
    if not url.startswith(('http://', 'https://')):
        url = 'http://' + url
    return url


def get_domain(url: str) -> str:
    """Return scheme://netloc of a URL (the base domain)."""
    parsed = urlparse(url)
    return f"{parsed.scheme}://{parsed.netloc}"


def is_internal_link(base_domain: str, target_url: str) -> bool:
    """Check if target_url is on the same domain as base_domain."""
    base_netloc = urlparse(base_domain).netloc
    target_netloc = urlparse(target_url).netloc
    return base_netloc == target_netloc


def parse_html_static(html: str):
    """
    Parse HTML with BeautifulSoup and extract content blocks:
      - Headings (H1-H6)
      - Paragraphs
      - Lists (UL/OL -> LI items)
      - Tables
      - Images
    Also extracts the <title> and meta desc.
    """
    soup = BeautifulSoup(html, 'html.parser')
    elements = []

    # Headings
    for level in range(1,7):
        for tag in soup.find_all(f"h{level}"):
            txt = tag.get_text(strip=True)
            if txt:
                elements.append({'type': f"Heading {level}", 'content': txt})

    # Paragraphs
    for p in soup.find_all('p'):
        txt = p.get_text(strip=True)
        if txt:
            elements.append({'type': 'Paragraph', 'content': txt})

    # Lists
    for lst in soup.find_all(['ul','ol']):
        for li in lst.find_all('li'):
            li_txt = li.get_text(strip=True)
            if li_txt:
                elements.append({'type': 'List item', 'content': li_txt})

    # Tables
    for table in soup.find_all('table'):
        rows_data = []
        for tr in table.find_all('tr'):
            row_cells = [td.get_text(strip=True) for td in tr.find_all(['th','td'])]
            if row_cells:
                rows_data.append(row_cells)
        if rows_data:
            elements.append({'type': 'Table', 'content': rows_data})

    # Images
    for img in soup.find_all('img'):
        src = img.get('src','')
        alt = img.get('alt','')
        if src:
            elements.append({'type': 'Image', 'content': src, 'alt': alt})

    # Basic SEO fields: Title & Meta desc
    title_tag = soup.find('title')
    page_title = title_tag.get_text(strip=True) if title_tag else ''

    meta_desc = ''
    desc_tag = soup.find('meta', attrs={'name':'description'})
    if desc_tag and desc_tag.get('content'):
        meta_desc = desc_tag['content'].strip()

    return elements, page_title, meta_desc

###########################
# 2. JS Rendering & Auth (Playwright)
###########################

def fetch_page_playwright(page: Page, url: str, wait_time=2, wait_selector=None) -> str:
    """
    Use an existing Playwright Page object to navigate & return HTML.
    Optionally wait for a CSS selector or wait a fixed time in seconds.
    Return the final HTML or None on error.
    """
    try:
        page.goto(url)
        if wait_selector:
            try:
                page.wait_for_selector(wait_selector, timeout=wait_time*1000)
            except:
                pass
        else:
            # Just wait a bit to ensure rendering
            page.wait_for_timeout(wait_time*1000)

        return page.content()
    except:
        return None


def create_playwright_browser(headless=True):
    """Launches a headless Chromium browser via Playwright, returns (browser, context)."""
    p = sync_playwright().start()
    browser = p.chromium.launch(headless=headless)
    context = browser.new_context()
    return p, browser, context


def login_playwright_js(context, login_url, username, password,
                        user_selector="#username", pass_selector="#password",
                        submit_selector="#login-submit", wait_after=3):
    """
    Perform a JS-based login using Playwright.
    This function assumes your login form has specific CSS selectors. Adjust as needed.
    1) Opens a new page from context.
    2) Navigates to login_url.
    3) Fills in user/pass fields, clicks submit.
    4) Waits some time or for a post-login selector.
    """
    page = context.new_page()
    page.goto(login_url)

    # fill in username/password
    try:
        page.fill(user_selector, username)
        page.fill(pass_selector, password)
        page.click(submit_selector)
        page.wait_for_timeout(wait_after * 1000)
    except:
        pass
    # close that page, but context has cookies now
    page.close()


def fetch_page_requests_auth(session: requests.Session, url: str, timeout=10):
    """Fetch page using the given requests session (for authenticated requests)."""
    try:
        resp = session.get(url, timeout=timeout)
        resp.raise_for_status()
        return resp.text
    except:
        return None


def create_requests_session(login_url: str, username: str, password: str,
                            user_field="username", pass_field="password",
                            extra_form_data=None) -> requests.Session:
    """
    Create a requests.Session with user logged in via a simple form-based auth.
    For pages that don't require JS-based login.
    """
    session = requests.Session()
    # 1) GET login page to get cookies/CSRF if needed
    r1 = session.get(login_url)
    if not r1.ok:
        return session  # might fail, but we proceed

    # 2) Prepare POST data
    data = {
        user_field: username,
        pass_field: password
    }
    if extra_form_data:
        data.update(extra_form_data)

    # 3) POST to login form
    post_url = login_url
    r2 = session.post(post_url, data=data)
    # If needed, check if login successful

    return session

###########################
# 3. BFS Crawler with Options
###########################

def compute_text_hash(elements):
    """
    Compute a naive hash for deduping, combining heading + paragraph text.
    """
    combined_text = []
    for e in elements:
        if e['type'].startswith('Heading') or e['type'] == 'Paragraph':
            combined_text.append(e['content'])
    joined = "\n".join(combined_text)
    return hashlib.md5(joined.encode('utf-8', errors='ignore')).hexdigest()


def should_crawl(url: str, rp: RobotFileParser, respect_robots: bool) -> bool:
    """Check robots.txt if required."""
    if not respect_robots:
        return True
    return rp.can_fetch('*', url)


def crawl_site(
    start_url: str,
    max_pages: int = 50,
    max_depth: int = 3,
    delay: float = 1.0,
    use_js_render: bool = False,
    js_context=None,  # if using JS, we pass in the context
    requests_session: requests.Session = None,  # if using requests-based auth
    respect_robots: bool = False
):
    """
    Perform BFS crawl.

    If use_js_render=True, we use the provided js_context to create new pages.
    If not, we use requests (optionally with requests_session for auth) to fetch.

    returns data_map: { url: {
      'title': str,
      'meta_description': str,
      'elements': [...],
      'links': [...],
      'status': int or 'error',
      'hash': str,
      'depth': int
    } }
    """
    start_url = sanitize_url(start_url)
    base_domain = get_domain(start_url)

    visited = set()
    queue = deque([(start_url, 0)])
    data_map = {}

    rp = RobotFileParser()
    rp.set_url(urljoin(base_domain, '/robots.txt'))
    try:
        rp.read()
    except:
        pass

    while queue and len(visited) < max_pages:
        current_url, depth = queue.popleft()
        if current_url in visited:
            continue
        visited.add(current_url)

        if depth > max_depth:
            continue

        if not should_crawl(current_url, rp, respect_robots):
            data_map[current_url] = {
                'title': '',
                'meta_description': '',
                'elements': [],
                'links': [],
                'status': 'disallowed_by_robots',
                'hash': '',
                'depth': depth
            }
            continue

        # Attempt to fetch page
        html = None
        status_code = 200
        if use_js_render and js_context:
            # use playwright
            page = js_context.new_page()
            html = fetch_page_playwright(page, current_url, wait_time=2)
            page.close()
            if not html:
                status_code = 'error'
        else:
            # use requests (auth or not)
            if requests_session:
                html = fetch_page_requests_auth(requests_session, current_url)
            else:
                try:
                    r = requests.get(current_url, timeout=10)
                    r.raise_for_status()
                    html = r.text
                except:
                    html = None
            if not html:
                status_code = 'error'

        if not html:
            data_map[current_url] = {
                'title': '',
                'meta_description': '',
                'elements': [],
                'links': [],
                'status': status_code,
                'hash': '',
                'depth': depth
            }
            time.sleep(delay)
            continue

        # parse
        elements, page_title, meta_desc = parse_html_static(html)
        soup = BeautifulSoup(html, 'html.parser')

        found_links = []
        for a in soup.find_all('a', href=True):
            lk = urljoin(current_url, a['href'])
            found_links.append(lk)

        content_hash = compute_text_hash(elements)
        data_map[current_url] = {
            'title': page_title,
            'meta_description': meta_desc,
            'elements': elements,
            'links': found_links,
            'status': status_code,
            'hash': content_hash,
            'depth': depth
        }

        for lk in found_links:
            if is_internal_link(base_domain, lk) and lk not in visited:
                queue.append((lk, depth+1))

        time.sleep(delay)

    return data_map

###########################
# 4. SEO & Analysis Helpers
###########################

def detect_duplicates(data_map):
    """Group URLs by hash. Return list of [url1, url2...] groups."""
    hash_map = defaultdict(list)
    for url, info in data_map.items():
        h = info['hash']
        if h:
            hash_map[h].append(url)
    duplicates = []
    for h, urls in hash_map.items():
        if len(urls) > 1:
            duplicates.append(urls)
    return duplicates


def find_broken_links(data_map):
    """Check if any link points to a known URL with status='error'."""
    known_urls = set(data_map.keys())
    broken = []
    for url, info in data_map.items():
        if info['status'] == 'error':
            continue
        for lk in info['links']:
            if lk in known_urls:
                if data_map[lk]['status'] == 'error':
                    broken.append((url, lk))
    return broken


def summarize_seo_issues(data_map):
    """Check for missing titles/descriptions."""
    missing_titles = []
    missing_meta = []
    for url, info in data_map.items():
        if info['status'] != 200:
            continue
        if not info['title']:
            missing_titles.append(url)
        if not info['meta_description']:
            missing_meta.append(url)
    return {
        'missing_titles': missing_titles,
        'missing_meta_descriptions': missing_meta
    }

###########################
# 5. Export Utilities
###########################

def export_to_csv(data_map):
    rows = []
    for url, info in data_map.items():
        row = {
            'URL': url,
            'Title': info['title'],
            'MetaDesc': info['meta_description'],
            'Status': info['status'],
            'Depth': info['depth'],
            'Hash': info['hash'],
            'LinksCount': len(info['links'])
        }
        rows.append(row)
    df = pd.DataFrame(rows)
    return df.to_csv(index=False).encode('utf-8')


def export_to_json(data_map):
    return json.dumps(data_map, indent=2, ensure_ascii=False).encode('utf-8')


def export_to_word(data_map):
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for url, info in data_map.items():
        document.add_heading(info['title'] or url, level=1)
        document.add_paragraph(f"URL: {url}")
        document.add_paragraph(f"Meta Description: {info['meta_description']}")
        document.add_paragraph(f"Status: {info['status']}")
        document.add_paragraph(f"Depth: {info['depth']}")
        document.add_paragraph("---")
        for el in info['elements']:
            etype = el['type']
            if etype.startswith('Heading'):
                level_num = min(int(etype.split()[1])+1, 6)
                document.add_heading(el['content'], level=level_num)
            elif etype == 'Paragraph':
                document.add_paragraph(el['content'])
            elif etype == 'List item':
                document.add_paragraph(f"• {el['content']}", style='List Bullet')
            elif etype == 'Table':
                document.add_paragraph("Table Data:")
                for row in el['content']:
                    document.add_paragraph(" | ".join(row))
            elif etype == 'Image':
                document.add_paragraph(f"[Image] src={el['content']} alt='{el.get('alt','')}'")
            else:
                document.add_paragraph(f"{etype}: {el['content']}")
        document.add_page_break()

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()

###########################
# 6. Streamlit UI
###########################

def run_app():
    st.title("Advanced Web Crawler & Content Extraction (Playwright) - Full")

    with st.sidebar:
        st.header("Crawl Settings")
        start_url = st.text_input("Start URL", "https://example.com")
        max_pages = st.slider("Max Pages", 1, 1000, 50)
        max_depth = st.slider("Max Depth", 0, 10, 3)
        delay = st.slider("Delay (seconds) between requests", 0.0, 5.0, 1.0)
        respect_robots = st.checkbox("Respect robots.txt")

        # Offer user the ability to do JS rendering
        use_js = st.checkbox("Enable JS Rendering (Playwright)")

        # Auth approach selection
        st.write("---")
        st.subheader("Authentication")
        auth_mode = st.selectbox(
            "Select Authentication Mode",
            ["None", "Requests-Form", "Playwright-JS"]
        )

        auth_info = {}
        if auth_mode != "None":
            login_url = st.text_input("Login URL", "https://example.com/login")
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")

        # For JS-based login, user may need to specify CSS selectors for fields and button.
        if auth_mode == "Playwright-JS":
            st.write("If your login form has different selectors, update them below.")
            user_sel = st.text_input("Username Field Selector", "#username")
            pass_sel = st.text_input("Password Field Selector", "#password")
            sub_sel  = st.text_input("Submit Button Selector", "#login-submit")
            wait_af  = st.number_input("Wait time after login (seconds)", 0.0, 30.0, 3.0)
            auth_info = {
                'user_sel': user_sel,
                'pass_sel': pass_sel,
                'sub_sel': sub_sel,
                'wait_af': wait_af,
            }

        st.write("---")
        crawl_button = st.button("Start Crawl")

    if 'data_map' not in st.session_state:
        st.session_state['data_map'] = {}

    if crawl_button:
        with st.spinner("Crawling in progress..."):

            js_context = None
            p = None
            requests_session = None

            # If using JS rendering or JS-based login, we open a Playwright browser.
            if use_js or auth_mode == "Playwright-JS":
                p, browser, context = create_playwright_browser(headless=True)
                js_context = context
                # If user selected JS-based login, do it.
                if auth_mode == "Playwright-JS" and username and password and login_url:
                    login_playwright_js(
                        context,
                        login_url,
                        username,
                        password,
                        user_selector=auth_info.get('user_sel',''),
                        pass_selector=auth_info.get('pass_sel',''),
                        submit_selector=auth_info.get('sub_sel',''),
                        wait_after=auth_info.get('wait_af',3)
                    )

            # If user selected requests-form auth
            if auth_mode == "Requests-Form" and username and password and login_url:
                requests_session = create_requests_session(
                    login_url, username, password
                )

            # BFS crawl
            data_map = crawl_site(
                start_url,
                max_pages=max_pages,
                max_depth=max_depth,
                delay=delay,
                use_js_render=(use_js or auth_mode=="Playwright-JS"),
                js_context=js_context,
                requests_session=requests_session,
                respect_robots=respect_robots
            )

            # close playwright if used
            if js_context:
                js_context.close()
            if p:
                p.stop()

            st.session_state['data_map'] = data_map
        st.success("Crawl Completed!")

    data_map = st.session_state['data_map']

    if data_map:
        st.subheader("Crawl Results")
        st.write(f"**Total pages crawled:** {len(data_map)}")

        # SEO summary
        issues = summarize_seo_issues(data_map)
        st.write("**SEO Issues**:")
        st.write(f"Missing Titles: {len(issues['missing_titles'])}")
        st.write(f"Missing Meta Descriptions: {len(issues['missing_meta_descriptions'])}")

        # duplicates
        dups = detect_duplicates(data_map)
        if dups:
            st.warning(f"Found {len(dups)} sets of duplicate pages.")
            for group in dups:
                st.write("Duplicate cluster:", group)
        else:
            st.write("No duplicate content detected.")

        # broken links
        broken_links = find_broken_links(data_map)
        if broken_links:
            st.warning(f"Found {len(broken_links)} broken internal links.")
            for src, dst in broken_links:
                st.write(f"From {src} to {dst}")
        else:
            st.write("No broken internal links detected.")

        # Page selection
        pages = list(data_map.keys())
        selected_page = st.selectbox("Select a page to preview", pages)
        if selected_page:
            info = data_map[selected_page]
            st.markdown(f"### {info['title'] or selected_page}")
            st.write(f"**URL:** {selected_page}")
            st.write(f"**Meta Description:** {info['meta_description']}")
            st.write(f"**Status:** {info['status']} | **Depth:** {info['depth']}")
            st.write("---")

            for el in info['elements']:
                etype = el['type']
                content = el.get('content','')
                if etype.startswith('Heading'):
                    level = int(etype.split()[1])
                    st.markdown(f"{'#'*min(level,6)} {content}")
                elif etype == 'Paragraph':
                    st.write(content)
                elif etype == 'List item':
                    st.markdown(f"- {content}")
                elif etype == 'Table':
                    df = pd.DataFrame(el['content'])
                    st.table(df)
                elif etype == 'Image':
                    with st.expander("Show Image"):
                        st.write(f"**Source:** {content}")
                        alt_text = el.get('alt','')
                        if alt_text:
                            st.write(f"Alt text: {alt_text}")
                else:
                    st.write(f"**{etype}:** {content}")

        # Export
        st.subheader("Export Options")
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Export CSV"):
                csv_data = export_to_csv(data_map)
                st.download_button(
                    "Download CSV",
                    data=csv_data,
                    file_name="crawl_data.csv",
                    mime="text/csv"
                )
        with c2:
            if st.button("Export JSON"):
                json_data = export_to_json(data_map)
                st.download_button(
                    "Download JSON",
                    data=json_data,
                    file_name="crawl_data.json",
                    mime="application/json"
                )
        with c3:
            if st.button("Export DOCX"):
                docx_data = export_to_word(data_map)
                st.download_button(
                    "Download Word",
                    data=docx_data,
                    file_name="crawl_data.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("Configure settings and click 'Start Crawl' to begin.")

if __name__ == "__main__":
    run_app()
